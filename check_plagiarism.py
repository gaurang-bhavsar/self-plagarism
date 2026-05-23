#!/usr/bin/env python3
"""
check_plagiarism.py - Self-Plagiarism Detection Tool

This script detects self-plagiarism (duplicate or highly similar content)
within a single document. It supports TXT, DOCX, and PDF formats.

Usage Examples:
    python check_plagiarism.py my_thesis.pdf
    python check_plagiarism.py essay.docx --output report.json --threshold 75
    python check_plagiarism.py document.txt -u paragraph -t 85

Technical Details:
    - Parses text, Word, and PDF documents.
    - Segments the document into sentences or paragraphs.
    - Performs exact and near-exact string similarity (word-set Jaccard + difflib).
    - Performs semantic similarity (TF-IDF cosine similarity, with sentence-transformers fallback).
    - Computes an overall duplication score based on word counts of flagged segments.
    - Outputs a detailed report (terminal/JSON/TXT) with severity levels.
"""

import os
import re
import sys
import json
import difflib
import argparse

# Optional and conditional imports
try:
    import docx
except ImportError:
    docx = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import nltk
except ImportError:
    nltk = None

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    TfidfVectorizer = None
    cosine_similarity = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

# ANSI colors for terminal output
COLOR_RESET = "\033[0m"
COLOR_BOLD = "\033[1m"
COLOR_RED = "\033[91m"
COLOR_GREEN = "\033[92m"
COLOR_YELLOW = "\033[93m"
COLOR_BLUE = "\033[94m"
COLOR_CYAN = "\033[96m"


def format_color(text, color_code, use_color=True):
    """Wraps text in ANSI color codes if color is enabled."""
    if use_color:
        return f"{color_code}{text}{COLOR_RESET}"
    return text


def read_text_file(path):
    """Reads a text file trying multiple common encodings."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16', 'utf-16le', 'utf-16be']
    for enc in encodings:
        try:
            with open(path, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise UnicodeError(f"Unable to decode text file '{path}' using standard encodings: {encodings}")


def read_docx_file(path):
    """Extracts text from a DOCX file, including paragraphs and tables."""
    if docx is None:
        raise ImportError("The 'python-docx' library is required to parse .docx files but is not installed.")
    
    try:
        doc = docx.Document(path)
    except Exception as e:
        raise ValueError(f"Failed to read DOCX file: {e}")
    
    content = []
    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            content.append(text)
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    content.append(text)
                    
    return "\n\n".join(content)


def read_pdf_file(path):
    """Extracts text from a PDF file using pdfplumber."""
    if pdfplumber is None:
        raise ImportError("The 'pdfplumber' library is required to parse .pdf files but is not installed.")
    
    try:
        content = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content.append(text.strip())
        return "\n\n".join(content)
    except Exception as e:
        raise ValueError(f"Failed to read PDF file: {e}")


def extract_text(path):
    """Extracts all text from a file based on its file extension."""
    if not os.path.exists(path):
        raise FileNotFoundError(f"File not found: '{path}'")
        
    _, ext = os.path.splitext(path.lower())
    if ext == '.txt':
        return read_text_file(path)
    elif ext == '.docx':
        return read_docx_file(path)
    elif ext == '.pdf':
        return read_pdf_file(path)
    else:
        # Fallback to text file reading
        try:
            return read_text_file(path)
        except Exception:
            raise ValueError(
                f"Unsupported file format: {ext}. "
                "Supported formats are .txt, .docx, and .pdf"
            )


def segment_text(text, unit='sentence'):
    """Segments the input text into a list of units (sentences or paragraphs)."""
    # Standardize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    if unit == 'paragraph':
        # Split by one or more blank lines
        raw_paras = re.split(r'\n\s*\n+', text)
        segments = []
        idx = 1
        for p in raw_paras:
            p_clean = re.sub(r'\s+', ' ', p).strip()
            if p_clean:
                segments.append({
                    'index': idx,
                    'text': p_clean,
                    'label': f"Paragraph {idx}"
                })
                idx += 1
        return segments
    else:  # sentence
        # Clean overall spacing first but preserve sentence structures
        text_clean = re.sub(r'\s+', ' ', text).strip()
        if not text_clean:
            return []
        
        sentences = []
        if nltk:
            try:
                # Ensure NLTK punkt tokenizer is available
                try:
                    nltk.data.find('tokenizers/punkt_tab')
                except LookupError:
                    nltk.download('punkt_tab', quiet=True)
                sentences = nltk.sent_tokenize(text_clean)
            except Exception:
                # Fallback to regex sentence tokenizer if NLTK fails
                sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s+', text_clean)
        else:
            # Fallback to regex sentence tokenizer
            sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s+', text_clean)
            
        segments = []
        idx = 1
        for s in sentences:
            s_clean = s.strip()
            if s_clean:
                segments.append({
                    'index': idx,
                    'text': s_clean,
                    'label': f"Sentence {idx}"
                })
                idx += 1
        return segments


def calculate_similarities(segments, threshold_pct, use_semantic_model=False, model_name=None):
    """
    Compares every segment against every other segment.
    Flags pairs with similarity >= threshold.
    """
    threshold = threshold_pct / 100.0
    n_segments = len(segments)
    
    if n_segments <= 1:
        return [], 0.0, set()
    
    texts = [s['text'] for s in segments]
    semantic_matrix = None
    
    # 1. Semantic Embedding / Matrix Computation
    if use_semantic_model and SentenceTransformer is not None:
        try:
            model = SentenceTransformer(model_name or 'all-MiniLM-L6-v2')
            embeddings = model.encode(texts, show_progress_bar=False)
            semantic_matrix = cosine_similarity(embeddings)
        except Exception as e:
            print(f"Warning: Failed to load sentence-transformers model. Falling back to TF-IDF. Error: {e}", file=sys.stderr)
            use_semantic_model = False
            
    if not use_semantic_model or semantic_matrix is None:
        if TfidfVectorizer is not None and cosine_similarity is not None:
            try:
                vectorizer = TfidfVectorizer(stop_words='english', token_pattern=r'(?u)\b\w+\b')
                tfidf_matrix = vectorizer.fit_transform(texts)
                semantic_matrix = cosine_similarity(tfidf_matrix)
                if hasattr(semantic_matrix, 'toarray'):
                    semantic_matrix = semantic_matrix.toarray()
            except ValueError:
                # Vocabulary empty (e.g. only stop words). Fallback to standard tf-idf without stop words
                try:
                    vectorizer = TfidfVectorizer(token_pattern=r'(?u)\b\w+\b')
                    tfidf_matrix = vectorizer.fit_transform(texts)
                    semantic_matrix = cosine_similarity(tfidf_matrix)
                    if hasattr(semantic_matrix, 'toarray'):
                        semantic_matrix = semantic_matrix.toarray()
                except Exception:
                    semantic_matrix = None
            except Exception:
                semantic_matrix = None

    # Pre-calculate lowercase clean word sets and lengths
    word_sets = []
    word_counts = []
    for s in segments:
        clean = re.sub(r'[^\w\s]', '', s['text'].lower())
        words = clean.split()
        word_sets.append(set(words))
        word_counts.append(len(words))

    flagged_pairs = []
    duplicate_indices = set()  # Indices of segments that are flagged as duplicate
    
    for i in range(n_segments):
        for j in range(i + 1, n_segments):
            text_i = texts[i]
            text_j = texts[j]
            
            # --- Exact / Near-exact Similarity ---
            w_i = word_sets[i]
            w_j = word_sets[j]
            
            # Word Jaccard
            union_w = w_i.union(w_j)
            jaccard = len(w_i.intersection(w_j)) / len(union_w) if union_w else 0.0
            
            clean_i = re.sub(r'\s+', '', text_i.lower())
            clean_j = re.sub(r'\s+', '', text_j.lower())
            
            if clean_i == clean_j:
                near_exact_score = 1.0
            elif jaccard >= 0.25 or (len(w_i.intersection(w_j)) >= 1 and (len(w_i) <= 3 or len(w_j) <= 3)):
                # Run difflib SequenceMatcher ratio for near-exact matching
                near_exact_score = difflib.SequenceMatcher(None, clean_i, clean_j).ratio()
            else:
                near_exact_score = 0.0
                
            # --- Semantic Similarity ---
            if semantic_matrix is not None:
                semantic_score = float(semantic_matrix[i][j])
                semantic_score = max(0.0, min(1.0, semantic_score))
            else:
                # If TF-IDF is not available, use Jaccard similarity as semantic proxy
                semantic_score = jaccard
                
            combined_score = max(near_exact_score, semantic_score)
            
            if combined_score >= threshold:
                flagged_pairs.append({
                    'index1': segments[i]['index'],
                    'label1': segments[i]['label'],
                    'text1': text_i,
                    'index2': segments[j]['index'],
                    'label2': segments[j]['label'],
                    'text2': text_j,
                    'near_exact_score': near_exact_score,
                    'semantic_score': semantic_score,
                    'combined_score': combined_score
                })
                # Mark segment j as a duplicate (later occurrence is the plagiarism)
                duplicate_indices.add(j)

    # Plagiarism score: percentage of words in duplicate segments
    total_words = sum(word_counts)
    duplicate_words = sum(word_counts[idx] for idx in duplicate_indices)
    overall_score = duplicate_words / total_words if total_words > 0 else 0.0
    
    return flagged_pairs, overall_score, duplicate_indices


def truncate_text(text, max_len=100):
    """Truncates text to max_len, appending ellipses if needed."""
    if len(text) <= max_len:
        return text
    return text[:max_len - 3] + "..."


def get_severity(score):
    """Determines severity category based on score."""
    if score < 0.10:
        return "Low"
    elif score <= 0.25:
        return "Medium"
    else:
        return "High"


def generate_report_text(file_path, unit, threshold, overall_score, severity, flagged_pairs, total_segments, use_color=False):
    """Generates the formatted plain text report (optionally with colors)."""
    score_pct = overall_score * 100.0
    
    # Severity Formatting
    sev_color = COLOR_GREEN
    if severity == "Medium":
        sev_color = COLOR_YELLOW
    elif severity == "High":
        sev_color = COLOR_RED
        
    severity_str = format_color(severity, sev_color + COLOR_BOLD if use_color else sev_color, use_color)
    score_str = format_color(f"{score_pct:.1f}%", COLOR_RED + COLOR_BOLD if score_pct > 25 else (COLOR_YELLOW + COLOR_BOLD if score_pct > 10 else COLOR_GREEN), use_color)
    
    lines = []
    lines.append("=" * 72)
    lines.append(format_color("SELF-PLAGIARISM DETECTION REPORT", COLOR_CYAN + COLOR_BOLD, use_color))
    lines.append("=" * 72)
    lines.append(f"File Path        : {file_path}")
    lines.append(f"Analysis Unit    : {unit.capitalize()}")
    lines.append(f"Total Units      : {total_segments}")
    lines.append(f"Match Threshold  : {threshold}%")
    lines.append("-" * 72)
    lines.append(f"Overall Score    : {score_str}")
    lines.append(f"Severity Rating  : {severity_str}")
    lines.append("=" * 72)
    lines.append("")
    
    if not flagged_pairs:
        lines.append(format_color("No duplicate segments detected above the threshold.", COLOR_GREEN, use_color))
        lines.append("")
        return "\n".join(lines)
        
    lines.append(format_color(f"Flagged Duplicate Pairs ({len(flagged_pairs)}):", COLOR_YELLOW + COLOR_BOLD, use_color))
    lines.append("")
    
    for i, pair in enumerate(flagged_pairs, 1):
        num_str = format_color(f"#{i}", COLOR_CYAN + COLOR_BOLD, use_color)
        lines.append(f"{num_str} {pair['label1']} vs {pair['label2']}")
        lines.append(f"  Position 1: {pair['label1']}")
        lines.append(f"  Content 1 : \"{truncate_text(pair['text1'])}\"")
        lines.append(f"  Position 2: {pair['label2']}")
        lines.append(f"  Content 2 : \"{truncate_text(pair['text2'])}\"")
        
        sim_val = pair['combined_score'] * 100
        ne_val = pair['near_exact_score'] * 100
        sem_val = pair['semantic_score'] * 100
        
        sim_str = format_color(f"{sim_val:.1f}%", COLOR_RED + COLOR_BOLD if sim_val >= 80 else COLOR_YELLOW, use_color)
        lines.append(f"  Similarity: {sim_str} [Near-Exact: {ne_val:.1f}%, Semantic: {sem_val:.1f}%]")
        lines.append("-" * 72)
        
    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(
        description="Detect self-plagiarism within a single document (TXT, DOCX, or PDF)."
    )
    parser.add_argument("file_path", help="Path to the document file to analyze.")
    parser.add_argument(
        "-t", "--threshold",
        type=float,
        default=80.0,
        help="Similarity threshold percentage (0-100). Default is 80."
    )
    parser.add_argument(
        "-u", "--unit",
        choices=["sentence", "paragraph"],
        default="sentence",
        help="Text unit for comparative analysis. Default is sentence."
    )
    parser.add_argument(
        "-o", "--output",
        help="Save report to this file (supported formats: .txt, .json)."
    )
    parser.add_argument(
        "--semantic-model",
        help="Use sentence-transformers with the specified model name (e.g. 'all-MiniLM-L6-v2')."
    )
    parser.add_argument(
        "--no-color",
        action="store_true",
        help="Disable ANSI color codes in terminal output."
    )
    
    args = parser.parse_args()
    
    # Enable color if stdout is a terminal and --no-color is not passed
    use_color = not args.no_color and sys.stdout.isatty()
    
    # Ensure threshold is in the range [0.0, 100.0]
    threshold = args.threshold
    if threshold < 0 or threshold > 100:
        print(f"Error: Threshold must be between 0 and 100. Got: {threshold}", file=sys.stderr)
        sys.exit(1)
        
    # Extract text from file
    try:
        text = extract_text(args.file_path)
    except FileNotFoundError as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"Error: Failed to read document: {e}", file=sys.stderr)
        sys.exit(1)
        
    # Segment text
    segments = segment_text(text, args.unit)
    
    if not segments:
        print(f"Error: Document '{args.file_path}' contains no text or is empty.", file=sys.stderr)
        sys.exit(1)
        
    # If the document only contains one segment, comparative analysis is not possible
    if len(segments) <= 1:
        overall_score = 0.0
        severity = "Low"
        flagged_pairs = []
        report_text = generate_report_text(
            args.file_path, args.unit, threshold, overall_score, severity, flagged_pairs, len(segments), use_color
        )
        print(report_text)
        print(format_color("Notice: Document contains 1 or fewer units. No comparative analysis can be performed.", COLOR_YELLOW, use_color))
        
        # Save to output file if specified
        if args.output:
            try:
                save_report(args.output, args.file_path, args.unit, threshold, overall_score, severity, flagged_pairs, len(segments))
            except Exception as e:
                print(f"Error saving output report: {e}", file=sys.stderr)
        return

    # Run plagiarism detection
    use_semantic_model = bool(args.semantic_model)
    try:
        flagged_pairs, overall_score, duplicate_indices = calculate_similarities(
            segments, threshold, use_semantic_model, args.semantic_model
        )
    except Exception as e:
        print(f"Error during similarity analysis: {e}", file=sys.stderr)
        sys.exit(1)
        
    severity = get_severity(overall_score)
    
    # Generate terminal report
    report_text = generate_report_text(
        args.file_path, args.unit, threshold, overall_score, severity, flagged_pairs, len(segments), use_color
    )
    print(report_text)
    
    # Save report to file if output is specified
    if args.output:
        try:
            save_report(args.output, args.file_path, args.unit, threshold, overall_score, severity, flagged_pairs, len(segments))
            print(f"Report saved to '{args.output}'")
        except Exception as e:
            print(f"Error: Failed to save report to '{args.output}': {e}", file=sys.stderr)


def save_report(output_path, file_path, unit, threshold, overall_score, severity, flagged_pairs, total_segments):
    """Saves the plagiarism report in TXT or JSON format."""
    _, ext = os.path.splitext(output_path.lower())
    
    if ext == '.json':
        report_data = {
            "file_path": file_path,
            "analysis_unit": unit,
            "total_segments": total_segments,
            "threshold_pct": threshold,
            "overall_score_pct": overall_score * 100.0,
            "severity_rating": severity,
            "flagged_pairs_count": len(flagged_pairs),
            "flagged_pairs": []
        }
        for pair in flagged_pairs:
            report_data["flagged_pairs"].append({
                "segment1": {
                    "index": pair["index1"],
                    "label": pair["label1"],
                    "text": pair["text1"]
                },
                "segment2": {
                    "index": pair["index2"],
                    "label": pair["label2"],
                    "text": pair["text2"]
                },
                "near_exact_score": pair["near_exact_score"],
                "semantic_score": pair["semantic_score"],
                "similarity_score": pair["combined_score"]
            })
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, indent=4)
    else:
        # Standard TXT report
        report_text = generate_report_text(
            file_path, unit, threshold, overall_score, severity, flagged_pairs, total_segments, use_color=False
        )
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(report_text)


if __name__ == "__main__":
    main()
